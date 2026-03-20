"""
Colass controller: resource/work catalog + estimator operations.
"""
from __future__ import annotations

import imaplib
import io
import json
import re
import unicodedata
from datetime import datetime
from email import message_from_bytes
from email.header import decode_header
from email.utils import parseaddr, parsedate_to_datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from config import Config
from models.database import DatabaseModel

try:
    from rapidfuzz import fuzz
except Exception:
    fuzz = None

try:
    from docx import Document
except Exception:
    Document = None

_WORD_CANONICAL = {
    # RU/RO common domain terms
    "truba": "teava",
    "труба": "teava",
    "teava": "teava",
    "țeava": "teava",
    "tevi": "teava",
    "трубы": "teava",
    "труб": "teava",
    "nisip": "nisip",
    "песок": "nisip",
    "pesok": "nisip",
    "sudura": "sudura",
    "sudură": "sudura",
    "sudor": "sudura",
    "сварка": "sudura",
    "сварщик": "sudura",
    "robinet": "robinet",
    "кран": "robinet",
    "kran": "robinet",
    "contor": "contor",
    "счетчик": "contor",
    "schetcik": "contor",
    "materiale": "materiale",
    "материалы": "materiale",
    "материал": "materiale",
    "manopera": "manopera",
    "работы": "manopera",
    "работа": "manopera",
    "utilaje": "utilaje",
    "техника": "utilaje",
    "оборудование": "utilaje",
    "gazoduct": "gazoduct",
    "газопровод": "gazoduct",
    "gaz": "gaz",
    "mufa": "mufa",
    "муфта": "mufa",
    "tranzitie": "tranzitie",
    "tranziție": "tranzitie",
    "переход": "tranzitie",
    "sedelka": "sedelka",
    "седелка": "sedelka",
    "cot": "cot",
    "отвод": "cot",
    "terasamente": "terasamente",
    "земляные": "terasamente",
    "sapatura": "sapatura",
    "копка": "sapatura",
}

_STOP_WORDS = {
    "si", "și", "и", "de", "cu", "la", "для", "на", "по", "în", "in", "din",
    "de", "al", "a", "the", "sau", "или", "pentru", "для", "об", "iz",
}


def _rows(result: Dict[str, Any], keys_lower: bool = True) -> List[Dict[str, Any]]:
    """Convert DatabaseModel.execute_query result to list[dict]."""
    if not result.get("success") or not result.get("columns"):
        return []
    cols = [c.upper() for c in (result.get("columns") or [])]
    out: List[Dict[str, Any]] = []
    for row in result.get("data") or []:
        d = dict(zip(cols, row))
        if keys_lower:
            d = {k.lower(): v for k, v in d.items() if k}
        out.append(d)
    return out


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        return float(value)
    except Exception:
        return default


def _norm_text(value: str) -> str:
    s = str(value or "").strip().lower()
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.replace("ё", "е")
    s = re.sub(r"[_/]+", " ", s)
    s = re.sub(r"[^\w\s.-]", " ", s, flags=re.UNICODE)
    s = re.sub(r"\s+", " ", s)
    words: List[str] = []
    for w in s.strip().split(" "):
        if not w:
            continue
        wl = _WORD_CANONICAL.get(w, w)
        if wl in _STOP_WORDS and len(wl) <= 3:
            continue
        words.append(wl)
    return " ".join(words).strip()


def _split_ai_parts(text: str) -> List[Dict[str, Any]]:
    if not text or not text.strip():
        return []
    normalized = re.sub(r"\s+(и|si|și)\s+", ", ", text, flags=re.IGNORECASE)
    chunks = [x.strip() for x in re.split(r"[,;\n]+", normalized) if x and x.strip()]
    out: List[Dict[str, Any]] = []
    for raw in chunks:
        qty = 1.0
        token = raw
        m = re.match(
            r"^\s*(\d+(?:[.,]\d+)?)\s*(?:шт|штук|buc|buc\.|m3|m2|m|kg|l|h-om|h-ut|ore|ч|час|часа|часов)?\s+(.+)$",
            raw,
            flags=re.IGNORECASE,
        )
        if m:
            qty = _safe_float((m.group(1) or "1").replace(",", "."), 1.0)
            token = m.group(2).strip()
        else:
            m2 = re.match(
                r"^(.+?)\s+(\d+(?:[.,]\d+)?)\s*(?:шт|штук|buc|buc\.|m3|m2|m|kg|l|h-om|h-ut|ore|ч|час|часа|часов)?\s*$",
                raw,
                flags=re.IGNORECASE,
            )
            if m2:
                token = m2.group(1).strip()
                qty = _safe_float((m2.group(2) or "1").replace(",", "."), 1.0)
        token = token.strip()
        if token:
            out.append({"original": raw, "token": token, "qty": max(qty, 0.0001)})
    return out


def _score_local(query_norm: str, name_norm: str, code_norm: str) -> float:
    if not query_norm or not name_norm:
        return 0.0
    if code_norm and query_norm == code_norm:
        return 100.0
    if query_norm == name_norm:
        return 99.0
    if query_norm in name_norm:
        return 92.0
    if name_norm in query_norm:
        return 80.0
    if fuzz is not None:
        return float(fuzz.WRatio(query_norm, name_norm))
    return 0.0


def _decode_mime_value(value: str) -> str:
    if not value:
        return ""
    out: List[str] = []
    for chunk, enc in decode_header(value):
        if isinstance(chunk, bytes):
            try:
                out.append(chunk.decode(enc or "utf-8", errors="replace"))
            except Exception:
                out.append(chunk.decode("utf-8", errors="replace"))
        else:
            out.append(str(chunk))
    return "".join(out).strip()


def _extract_email_text(msg: Any) -> str:
    if msg is None:
        return ""
    parts: List[str] = []
    if msg.is_multipart():
        for part in msg.walk():
            ctype = (part.get_content_type() or "").lower()
            disp = (part.get("Content-Disposition") or "").lower()
            if "attachment" in disp:
                continue
            if ctype in ("text/plain", "text/html"):
                payload = part.get_payload(decode=True)
                if payload is None:
                    continue
                charset = part.get_content_charset() or "utf-8"
                text = payload.decode(charset, errors="replace")
                if ctype == "text/html":
                    text = re.sub(r"<[^>]+>", " ", text)
                parts.append(text)
    else:
        payload = msg.get_payload(decode=True)
        if isinstance(payload, bytes):
            charset = msg.get_content_charset() or "utf-8"
            parts.append(payload.decode(charset, errors="replace"))
        elif isinstance(payload, str):
            parts.append(payload)
    text = "\n".join(parts)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def _guess_lang_pref(text: str) -> str:
    s = (text or "").lower()
    if re.search(r"[а-яё]", s):
        return "ru"
    ro_hits = 0
    for tok in ("ș", "ț", "ă", "î", "â", "deviz", "oferta", "resurse", "lucrari", "lucrare"):
        if tok in s:
            ro_hits += 1
    return "ro" if ro_hits else "ru"


def _parse_contact_from_text(text: str) -> Dict[str, Optional[str]]:
    s = text or ""
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", s)
    phone_match = re.search(r"(\+?\d[\d\-\s()]{6,}\d)", s)
    name_match = re.search(r"(?:контакт|contact|persoan[ăa]\s+de\s+contact)\s*[:\-]\s*([^\n,;]+)", s, flags=re.IGNORECASE)
    company_match = re.search(r"(?:компан|company|companie)\s*[:\-]\s*([^\n,;]+)", s, flags=re.IGNORECASE)
    return {
        "contact_name": (name_match.group(1).strip() if name_match else None),
        "company_name": (company_match.group(1).strip() if company_match else None),
        "phone": (phone_match.group(1).strip() if phone_match else None),
        "email": (email_match.group(0).strip() if email_match else None),
    }


class ColassController:
    """API for Colass module: catalog and estimator."""

    @staticmethod
    def _recalc_totals(db: DatabaseModel, estimate_id: int) -> None:
        """Recalculate estimate totals by resource types + section totals."""
        r = db.execute_query(
            """
            SELECT NVL(rt.CODE, 'OTHER') AS TYPE_CODE,
                   NVL(SUM(i.AMOUNT), 0) AS TOTAL_AMOUNT
            FROM CLS_ESTIMATE_ITEMS i
            LEFT JOIN CLS_RESOURCE_TYPES rt ON rt.ID = i.RESOURCE_TYPE_ID
            WHERE i.ESTIMATE_ID = :eid
            GROUP BY NVL(rt.CODE, 'OTHER')
            """,
            {"eid": estimate_id},
        )
        by_type = {x.get("type_code"): _safe_float(x.get("total_amount")) for x in _rows(r)}

        total_manopera = by_type.get("MANOPERA", 0.0)
        total_materiale = by_type.get("MATERIALE", 0.0)
        total_utilaje = by_type.get("UTILAJE", 0.0)
        total_amount = sum(by_type.values())

        db.execute_query(
            """
            UPDATE CLS_ESTIMATES
            SET TOTAL_MANOPERA = :m1,
                TOTAL_MATERIALE = :m2,
                TOTAL_UTILAJE = :m3,
                TOTAL_AMOUNT = :mt,
                UPDATED_AT = SYSTIMESTAMP
            WHERE ID = :eid
            """,
            {
                "eid": estimate_id,
                "m1": total_manopera,
                "m2": total_materiale,
                "m3": total_utilaje,
                "mt": total_amount,
            },
        )

        sec = db.execute_query(
            """
            SELECT SECTION_ID, NVL(SUM(AMOUNT), 0) AS TOTAL_AMOUNT
            FROM CLS_ESTIMATE_ITEMS
            WHERE ESTIMATE_ID = :eid AND SECTION_ID IS NOT NULL
            GROUP BY SECTION_ID
            """,
            {"eid": estimate_id},
        )
        sec_totals = {int(x["section_id"]): _safe_float(x["total_amount"]) for x in _rows(sec)}

        all_sections = db.execute_query(
            "SELECT ID FROM CLS_ESTIMATE_SECTIONS WHERE ESTIMATE_ID = :eid",
            {"eid": estimate_id},
        )
        for section in _rows(all_sections):
            sid = int(section["id"])
            db.execute_query(
                """
                UPDATE CLS_ESTIMATE_SECTIONS
                SET TOTAL_AMOUNT = :total,
                    UPDATED_AT = SYSTIMESTAMP
                WHERE ID = :sid
                """,
                {"sid": sid, "total": sec_totals.get(sid, 0.0)},
            )

    @staticmethod
    def get_catalog_tree(search: Optional[str] = None) -> Dict[str, Any]:
        """Hierarchy tree for catalog: level1 -> level2 -> level3 -> works."""
        try:
            with DatabaseModel() as db:
                sql = """
                    SELECT ID, WORK_NO, WORK_CODE, WORK_NAME, UNIT, PRICE, AMOUNT,
                           LEVEL1_NAME, LEVEL2_NAME, LEVEL3_NAME
                    FROM V_CLS_WORK_TREE
                    WHERE IS_ACTIVE = 'Y'
                """
                params: Dict[str, Any] = {}
                if search and search.strip():
                    sql += " AND (UPPER(WORK_NAME) LIKE '%' || UPPER(:q) || '%' OR UPPER(NVL(WORK_CODE,'')) LIKE '%' || UPPER(:q) || '%')"
                    params["q"] = search.strip()
                sql += " ORDER BY WORK_NO"
                rows = _rows(db.execute_query(sql, params))

            tree: Dict[str, Dict[str, Dict[str, List[Dict[str, Any]]]]] = {}
            for row in rows:
                l1 = (row.get("level1_name") or "Без раздела").strip()
                l2 = (row.get("level2_name") or "Без подраздела").strip()
                l3 = (row.get("level3_name") or "Без группы").strip()
                tree.setdefault(l1, {}).setdefault(l2, {}).setdefault(l3, []).append(
                    {
                        "id": row.get("id"),
                        "work_no": row.get("work_no"),
                        "work_code": row.get("work_code"),
                        "work_name": row.get("work_name"),
                        "unit": row.get("unit"),
                        "price": row.get("price"),
                        "amount": row.get("amount"),
                    }
                )

            out = []
            for l1_name, l2_map in tree.items():
                l1_children = []
                for l2_name, l3_map in l2_map.items():
                    l2_children = []
                    for l3_name, works in l3_map.items():
                        l2_children.append({"name": l3_name, "works": works})
                    l1_children.append({"name": l2_name, "children": l2_children})
                out.append({"name": l1_name, "children": l1_children})

            return {"success": True, "data": out, "count": len(rows)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_work_catalog(search: Optional[str] = None, limit: int = 500) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                sql = """
                    SELECT * FROM (
                        SELECT w.ID, w.WORK_NO, w.WORK_CODE, w.WORK_NAME, w.UNIT, w.PRICE, w.AMOUNT,
                               w.LEVEL1_NAME, w.LEVEL2_NAME, w.LEVEL3_NAME,
                               (SELECT COUNT(*) FROM CLS_WORK_RESOURCES wr WHERE wr.WORK_ID = w.ID) AS RESOURCES_COUNT
                        FROM V_CLS_WORK_TREE w
                        WHERE w.IS_ACTIVE = 'Y'
                """
                params: Dict[str, Any] = {"lim": max(1, min(limit, 2000))}
                if search and search.strip():
                    sql += " AND (UPPER(w.WORK_NAME) LIKE '%' || UPPER(:q) || '%' OR UPPER(NVL(w.WORK_CODE,'')) LIKE '%' || UPPER(:q) || '%')"
                    params["q"] = search.strip()
                sql += " ORDER BY w.WORK_NO ) WHERE ROWNUM <= :lim"
                data = _rows(db.execute_query(sql, params))
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_resources(
        search: Optional[str] = None,
        type_code: Optional[str] = None,
        limit: int = 1000,
    ) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                sql = """
                    SELECT * FROM (
                        SELECT ID, RESOURCE_CODE, NAME, UNIT, BASE_QTY, PRICE, AMOUNT,
                               TYPE_ID, TYPE_CODE, TYPE_NAME_RU, TYPE_NAME_RO, TYPE_NAME_EN,
                               SOURCE_DOC, SOURCE_ROW
                        FROM V_CLS_RESOURCE_CATALOG
                        WHERE 1=1
                """
                params: Dict[str, Any] = {"lim": max(1, min(limit, 5000))}
                if type_code and type_code.strip():
                    sql += " AND UPPER(TYPE_CODE) = UPPER(:type_code)"
                    params["type_code"] = type_code.strip()
                if search and search.strip():
                    sql += " AND (UPPER(NAME) LIKE '%' || UPPER(:q) || '%' OR UPPER(RESOURCE_CODE) LIKE '%' || UPPER(:q) || '%')"
                    params["q"] = search.strip()
                sql += " ORDER BY TYPE_ID, RESOURCE_CODE ) WHERE ROWNUM <= :lim"
                data = _rows(db.execute_query(sql, params))
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_work_resources(work_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                r = db.execute_query(
                    """
                    SELECT wr.ID, wr.WORK_ID, wr.RESOURCE_ID, wr.RESOURCE_CODE, wr.RESOURCE_NAME,
                           wr.UNIT, wr.QTY_NORM, wr.UNIT_PRICE, wr.AMOUNT, wr.SORT_ORDER,
                           rt.ID AS TYPE_ID, rt.CODE AS TYPE_CODE, rt.NAME_RU AS TYPE_NAME_RU, rt.NAME_RO AS TYPE_NAME_RO
                    FROM CLS_WORK_RESOURCES wr
                    LEFT JOIN CLS_RESOURCES r ON r.ID = wr.RESOURCE_ID
                    LEFT JOIN CLS_RESOURCE_TYPES rt ON rt.ID = r.TYPE_ID
                    WHERE wr.WORK_ID = :wid
                    ORDER BY wr.SORT_ORDER, wr.ID
                    """,
                    {"wid": work_id},
                )
                data = _rows(r)
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_projects() -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                r = db.execute_query(
                    """
                    SELECT ID, CODE, NAME, CUSTOMER_NAME, OBJECT_TYPE, LOCATION, CURRENCY, STATUS, CREATED_AT
                    FROM CLS_PROJECTS
                    ORDER BY ID DESC
                    """
                )
                data = _rows(r)
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_project_estimates(project_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                r = db.execute_query(
                    """
                    SELECT ID, PROJECT_ID, CODE, NAME, CURRENCY, STATUS,
                           TOTAL_MANOPERA, TOTAL_MATERIALE, TOTAL_UTILAJE, TOTAL_AMOUNT,
                           CREATED_AT, UPDATED_AT
                    FROM CLS_ESTIMATES
                    WHERE PROJECT_ID = :pid
                    ORDER BY ID DESC
                    """,
                    {"pid": project_id},
                )
                data = _rows(r)
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_estimate_detail(estimate_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                r = db.execute_query(
                    """
                    SELECT e.ID, e.PROJECT_ID, e.CODE, e.NAME, e.CURRENCY, e.STATUS,
                           e.TOTAL_MANOPERA, e.TOTAL_MATERIALE, e.TOTAL_UTILAJE, e.TOTAL_AMOUNT,
                           e.SOURCE_DOC, e.SOURCE_REF, e.CREATED_AT, e.UPDATED_AT,
                           p.NAME AS PROJECT_NAME
                    FROM CLS_ESTIMATES e
                    JOIN CLS_PROJECTS p ON p.ID = e.PROJECT_ID
                    WHERE e.ID = :eid
                    """,
                    {"eid": estimate_id},
                )
                rows = _rows(r)
                return {"success": True, "data": rows[0] if rows else None}
        except Exception as e:
            return {"success": False, "error": str(e), "data": None}

    @staticmethod
    def get_estimate_sections(estimate_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                r = db.execute_query(
                    """
                    SELECT ID, ESTIMATE_ID, CODE, NAME, SORT_ORDER, TOTAL_AMOUNT
                    FROM CLS_ESTIMATE_SECTIONS
                    WHERE ESTIMATE_ID = :eid
                    ORDER BY SORT_ORDER, ID
                    """,
                    {"eid": estimate_id},
                )
                data = _rows(r)
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_estimate_items(estimate_id: int, section_id: Optional[int] = None) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                sql = """
                    SELECT i.ID, i.ESTIMATE_ID, i.SECTION_ID, s.NAME AS SECTION_NAME,
                           i.WORK_ID, i.RESOURCE_TYPE_ID, i.RESOURCE_ID,
                           i.ITEM_CODE, i.ITEM_NAME, i.UNIT, i.QTY, i.PRICE, i.AMOUNT,
                           i.NOTES, i.SOURCE_DOC, i.SOURCE_ROW, i.CREATED_BY,
                           rt.CODE AS TYPE_CODE, rt.NAME_RU AS TYPE_NAME_RU, rt.NAME_RO AS TYPE_NAME_RO
                    FROM CLS_ESTIMATE_ITEMS i
                    LEFT JOIN CLS_ESTIMATE_SECTIONS s ON s.ID = i.SECTION_ID
                    LEFT JOIN CLS_RESOURCE_TYPES rt ON rt.ID = i.RESOURCE_TYPE_ID
                    WHERE i.ESTIMATE_ID = :eid
                """
                params: Dict[str, Any] = {"eid": estimate_id}
                if section_id is not None:
                    sql += " AND i.SECTION_ID = :sid"
                    params["sid"] = section_id
                sql += " ORDER BY NVL(s.SORT_ORDER, 999), i.ID"
                data = _rows(db.execute_query(sql, params))
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def add_estimate_item(estimate_id: int, payload: Dict[str, Any]) -> Dict[str, Any]:
        try:
            section_id = payload.get("section_id")
            work_id = payload.get("work_id")
            resource_type_id = payload.get("resource_type_id")
            resource_id = payload.get("resource_id")
            item_code = (payload.get("item_code") or "").strip() or None
            item_name = (payload.get("item_name") or "").strip()
            unit = (payload.get("unit") or "").strip() or None
            qty = _safe_float(payload.get("qty"), 0.0)
            price = _safe_float(payload.get("price"), 0.0)
            amount = _safe_float(payload.get("amount"), qty * price)
            notes = (payload.get("notes") or "").strip() or None
            created_by = (payload.get("created_by") or "ui").strip()[:120]
            source_doc = (payload.get("source_doc") or "MANUAL").strip()[:30]
            source_row = payload.get("source_row")

            if not item_name:
                return {"success": False, "error": "item_name required"}

            with DatabaseModel() as db:
                db.execute_query(
                    """
                    INSERT INTO CLS_ESTIMATE_ITEMS
                        (ESTIMATE_ID, SECTION_ID, WORK_ID, RESOURCE_TYPE_ID, RESOURCE_ID,
                         ITEM_CODE, ITEM_NAME, UNIT, QTY, PRICE, AMOUNT,
                         NOTES, SOURCE_DOC, SOURCE_ROW, CREATED_BY)
                    VALUES
                        (:estimate_id, :section_id, :work_id, :resource_type_id, :resource_id,
                         :item_code, :item_name, :unit, :qty, :price, :amount,
                         :notes, :source_doc, :source_row, :created_by)
                    """,
                    {
                        "estimate_id": estimate_id,
                        "section_id": section_id,
                        "work_id": work_id,
                        "resource_type_id": resource_type_id,
                        "resource_id": resource_id,
                        "item_code": item_code,
                        "item_name": item_name,
                        "unit": unit,
                        "qty": qty,
                        "price": price,
                        "amount": amount,
                        "notes": notes,
                        "source_doc": source_doc,
                        "source_row": source_row,
                        "created_by": created_by,
                    },
                )
                ColassController._recalc_totals(db, estimate_id)
                db.connection.commit()

                rid = db.execute_query(
                    "SELECT MAX(ID) AS ID FROM CLS_ESTIMATE_ITEMS WHERE ESTIMATE_ID = :eid",
                    {"eid": estimate_id},
                )
                row = _rows(rid)
                new_id = row[0].get("id") if row else None
                return {"success": True, "id": new_id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def update_estimate_item(item_id: int, payload: Dict[str, Any]) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                base = _rows(
                    db.execute_query(
                        "SELECT ID, ESTIMATE_ID, SECTION_ID, QTY, PRICE, AMOUNT, NOTES FROM CLS_ESTIMATE_ITEMS WHERE ID = :id",
                        {"id": item_id},
                    )
                )
                if not base:
                    return {"success": False, "error": "Item not found"}
                cur = base[0]

                section_id = payload.get("section_id", cur.get("section_id"))
                qty = _safe_float(payload.get("qty", cur.get("qty")), 0.0)
                price = _safe_float(payload.get("price", cur.get("price")), 0.0)
                amount = payload.get("amount")
                if amount is None:
                    amount = qty * price
                amount = _safe_float(amount, qty * price)
                notes = payload.get("notes", cur.get("notes"))

                db.execute_query(
                    """
                    UPDATE CLS_ESTIMATE_ITEMS
                    SET SECTION_ID = :section_id,
                        QTY = :qty,
                        PRICE = :price,
                        AMOUNT = :amount,
                        NOTES = :notes,
                        UPDATED_AT = SYSTIMESTAMP
                    WHERE ID = :id
                    """,
                    {
                        "id": item_id,
                        "section_id": section_id,
                        "qty": qty,
                        "price": price,
                        "amount": amount,
                        "notes": notes,
                    },
                )
                ColassController._recalc_totals(db, int(cur["estimate_id"]))
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def delete_estimate_item(item_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                item = _rows(
                    db.execute_query(
                        "SELECT ESTIMATE_ID FROM CLS_ESTIMATE_ITEMS WHERE ID = :id",
                        {"id": item_id},
                    )
                )
                if not item:
                    return {"success": False, "error": "Item not found"}
                estimate_id = int(item[0]["estimate_id"])

                db.execute_query("DELETE FROM CLS_ESTIMATE_ITEMS WHERE ID = :id", {"id": item_id})
                ColassController._recalc_totals(db, estimate_id)
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_estimate_summary(estimate_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                totals = _rows(
                    db.execute_query(
                        """
                        SELECT ID, CODE, NAME, TOTAL_MANOPERA, TOTAL_MATERIALE, TOTAL_UTILAJE, TOTAL_AMOUNT
                        FROM CLS_ESTIMATES
                        WHERE ID = :eid
                        """,
                        {"eid": estimate_id},
                    )
                )
                by_type = _rows(
                    db.execute_query(
                        """
                        SELECT NVL(rt.CODE, 'OTHER') AS TYPE_CODE,
                               NVL(rt.NAME_RU, 'Прочее') AS TYPE_NAME_RU,
                               NVL(rt.NAME_RO, 'Altele') AS TYPE_NAME_RO,
                               COUNT(i.ID) AS ITEMS_COUNT,
                               NVL(SUM(i.AMOUNT), 0) AS TOTAL_AMOUNT
                        FROM CLS_ESTIMATE_ITEMS i
                        LEFT JOIN CLS_RESOURCE_TYPES rt ON rt.ID = i.RESOURCE_TYPE_ID
                        WHERE i.ESTIMATE_ID = :eid
                        GROUP BY NVL(rt.CODE, 'OTHER'), NVL(rt.NAME_RU, 'Прочее'), NVL(rt.NAME_RO, 'Altele')
                        ORDER BY TYPE_CODE
                        """,
                        {"eid": estimate_id},
                    )
                )
                by_sections = _rows(
                    db.execute_query(
                        """
                        SELECT s.ID, s.CODE, s.NAME, s.SORT_ORDER, s.TOTAL_AMOUNT,
                               (SELECT COUNT(*) FROM CLS_ESTIMATE_ITEMS i WHERE i.SECTION_ID = s.ID) AS ITEMS_COUNT
                        FROM CLS_ESTIMATE_SECTIONS s
                        WHERE s.ESTIMATE_ID = :eid
                        ORDER BY s.SORT_ORDER, s.ID
                        """,
                        {"eid": estimate_id},
                    )
                )
                return {
                    "success": True,
                    "estimate": totals[0] if totals else None,
                    "by_type": by_type,
                    "by_sections": by_sections,
                }
        except Exception as e:
            return {"success": False, "error": str(e), "estimate": None, "by_type": [], "by_sections": []}

    @staticmethod
    def add_work_to_estimate(
        estimate_id: int,
        work_id: int,
        section_id: Optional[int] = None,
        multiplier: float = 1.0,
        created_by: str = "ui",
    ) -> Dict[str, Any]:
        """Add work resources to estimate as rows in CLS_ESTIMATE_ITEMS."""
        try:
            multiplier = max(0.0, _safe_float(multiplier, 1.0))
            with DatabaseModel() as db:
                work_rows = _rows(
                    db.execute_query(
                        """
                        SELECT ID, WORK_NO, WORK_CODE, WORK_NAME, UNIT, PRICE, AMOUNT
                        FROM CLS_WORK_CATALOG
                        WHERE ID = :id
                        """,
                        {"id": work_id},
                    )
                )
                if not work_rows:
                    return {"success": False, "error": "Work not found"}
                work = work_rows[0]

                if section_id is None:
                    first_section = _rows(
                        db.execute_query(
                            "SELECT ID FROM CLS_ESTIMATE_SECTIONS WHERE ESTIMATE_ID = :eid ORDER BY SORT_ORDER, ID",
                            {"eid": estimate_id},
                        )
                    )
                    section_id = first_section[0].get("id") if first_section else None

                res_rows = _rows(
                    db.execute_query(
                        """
                        SELECT wr.RESOURCE_ID, wr.RESOURCE_CODE, wr.RESOURCE_NAME,
                               wr.UNIT, wr.QTY_NORM, wr.UNIT_PRICE, wr.AMOUNT,
                               r.TYPE_ID AS RESOURCE_TYPE_ID
                        FROM CLS_WORK_RESOURCES wr
                        LEFT JOIN CLS_RESOURCES r ON r.ID = wr.RESOURCE_ID
                        WHERE wr.WORK_ID = :wid
                        ORDER BY wr.SORT_ORDER, wr.ID
                        """,
                        {"wid": work_id},
                    )
                )

                inserted = 0
                if res_rows:
                    for res in res_rows:
                        qty = _safe_float(res.get("qty_norm"), 0.0) * multiplier
                        price = _safe_float(res.get("unit_price"), 0.0)
                        amount = qty * price
                        if amount == 0:
                            amount = _safe_float(res.get("amount"), 0.0) * multiplier

                        db.execute_query(
                            """
                            INSERT INTO CLS_ESTIMATE_ITEMS
                                (ESTIMATE_ID, SECTION_ID, WORK_ID, RESOURCE_TYPE_ID, RESOURCE_ID,
                                 ITEM_CODE, ITEM_NAME, UNIT, QTY, PRICE, AMOUNT,
                                 SOURCE_DOC, CREATED_BY)
                            VALUES
                                (:estimate_id, :section_id, :work_id, :resource_type_id, :resource_id,
                                 :item_code, :item_name, :unit, :qty, :price, :amount,
                                 'F5', :created_by)
                            """,
                            {
                                "estimate_id": estimate_id,
                                "section_id": section_id,
                                "work_id": work_id,
                                "resource_type_id": res.get("resource_type_id"),
                                "resource_id": res.get("resource_id"),
                                "item_code": res.get("resource_code") or work.get("work_code"),
                                "item_name": res.get("resource_name") or work.get("work_name"),
                                "unit": res.get("unit") or work.get("unit"),
                                "qty": qty,
                                "price": price,
                                "amount": amount,
                                "created_by": created_by[:120],
                            },
                        )
                        inserted += 1
                else:
                    qty = max(multiplier, 1.0)
                    price = _safe_float(work.get("price"), 0.0)
                    amount = qty * price
                    if amount == 0:
                        amount = _safe_float(work.get("amount"), 0.0) * qty
                    db.execute_query(
                        """
                        INSERT INTO CLS_ESTIMATE_ITEMS
                            (ESTIMATE_ID, SECTION_ID, WORK_ID, ITEM_CODE, ITEM_NAME, UNIT, QTY, PRICE, AMOUNT, SOURCE_DOC, CREATED_BY)
                        VALUES
                            (:estimate_id, :section_id, :work_id, :item_code, :item_name, :unit, :qty, :price, :amount, 'F5', :created_by)
                        """,
                        {
                            "estimate_id": estimate_id,
                            "section_id": section_id,
                            "work_id": work_id,
                            "item_code": work.get("work_code"),
                            "item_name": work.get("work_name"),
                            "unit": work.get("unit"),
                            "qty": qty,
                            "price": price,
                            "amount": amount,
                            "created_by": created_by[:120],
                        },
                    )
                    inserted = 1

                ColassController._recalc_totals(db, estimate_id)
                db.connection.commit()
                return {"success": True, "inserted": inserted, "work_id": work_id, "estimate_id": estimate_id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def ai_parse_estimate_local(text: str, threshold: int = 45) -> Dict[str, Any]:
        """Local AI parse using rapidfuzz fallback scoring."""
        try:
            parts = _split_ai_parts(text)
            if not parts:
                return {"success": True, "matches": [], "backend": "local"}

            with DatabaseModel() as db:
                works = _rows(
                    db.execute_query(
                        """
                        SELECT ID, WORK_CODE AS CODE, WORK_NAME AS NAME, UNIT, NVL(PRICE, NVL(AMOUNT, 0)) AS PRICE
                        FROM CLS_WORK_CATALOG
                        WHERE IS_ACTIVE = 'Y'
                        """
                    )
                )
                resources = _rows(
                    db.execute_query(
                        """
                        SELECT ID, RESOURCE_CODE AS CODE, NAME, UNIT, NVL(PRICE, NVL(AMOUNT, 0)) AS PRICE, TYPE_ID
                        FROM CLS_RESOURCES
                        WHERE IS_ACTIVE = 'Y'
                        """
                    )
                )

            matches: List[Dict[str, Any]] = []
            for part in parts:
                token = part["token"]
                token_norm = _norm_text(token)
                qty = part["qty"]
                best: Optional[Dict[str, Any]] = None
                best_score = -1.0

                for row in works:
                    name_norm = _norm_text(row.get("name") or "")
                    code_norm = _norm_text(row.get("code") or "")
                    score = _score_local(token_norm, name_norm, code_norm)
                    if score > best_score:
                        best_score = score
                        best = {
                            "entity_type": "work",
                            "id": row.get("id"),
                            "code": row.get("code"),
                            "name": row.get("name"),
                            "unit": row.get("unit"),
                            "price": _safe_float(row.get("price"), 0.0),
                            "qty": qty,
                            "confidence": round(score, 2),
                            "original": part["original"],
                            "token": token,
                        }

                for row in resources:
                    name_norm = _norm_text(row.get("name") or "")
                    code_norm = _norm_text(row.get("code") or "")
                    score = _score_local(token_norm, name_norm, code_norm)
                    if score > best_score:
                        best_score = score
                        best = {
                            "entity_type": "resource",
                            "id": row.get("id"),
                            "code": row.get("code"),
                            "name": row.get("name"),
                            "unit": row.get("unit"),
                            "price": _safe_float(row.get("price"), 0.0),
                            "resource_type_id": row.get("type_id"),
                            "qty": qty,
                            "confidence": round(score, 2),
                            "original": part["original"],
                            "token": token,
                        }

                if best and best["confidence"] >= threshold:
                    matches.append(best)
                else:
                    matches.append(
                        {
                            "entity_type": None,
                            "id": None,
                            "code": None,
                            "name": None,
                            "unit": None,
                            "price": 0.0,
                            "qty": qty,
                            "confidence": round(max(best_score, 0), 2),
                            "original": part["original"],
                            "token": token,
                        }
                    )

            return {"success": True, "matches": matches, "backend": "local"}
        except Exception as e:
            return {"success": False, "error": str(e), "matches": [], "backend": "local"}

    @staticmethod
    def ai_parse_estimate_oracle(text: str, threshold: int = 45) -> Dict[str, Any]:
        """Oracle AI parse using UTL_MATCH similarity."""
        try:
            parts = _split_ai_parts(text)
            if not parts:
                return {"success": True, "matches": [], "backend": "oracle"}

            matches: List[Dict[str, Any]] = []
            with DatabaseModel() as db:
                for part in parts:
                    token = part["token"]
                    qty = part["qty"]
                    r = db.execute_query(
                        """
                        SELECT * FROM (
                            SELECT
                                'work' AS ENTITY_TYPE,
                                w.ID AS ENTITY_ID,
                                w.WORK_CODE AS CODE,
                                w.WORK_NAME AS NAME,
                                w.UNIT AS UNIT,
                                NVL(w.PRICE, NVL(w.AMOUNT, 0)) AS PRICE,
                                CAST(NULL AS NUMBER) AS RESOURCE_TYPE_ID,
                                GREATEST(
                                    CASE WHEN UPPER(NVL(w.WORK_CODE, '')) = UPPER(:q) THEN 100 ELSE 0 END,
                                    CASE WHEN UPPER(w.WORK_NAME) LIKE '%' || UPPER(:q) || '%' THEN 90 ELSE 0 END,
                                    UTL_MATCH.JARO_WINKLER_SIMILARITY(UPPER(w.WORK_NAME), UPPER(:q)),
                                    UTL_MATCH.EDIT_DISTANCE_SIMILARITY(UPPER(w.WORK_NAME), UPPER(:q))
                                ) AS SCORE
                            FROM CLS_WORK_CATALOG w
                            WHERE w.IS_ACTIVE = 'Y'

                            UNION ALL

                            SELECT
                                'resource' AS ENTITY_TYPE,
                                r.ID AS ENTITY_ID,
                                r.RESOURCE_CODE AS CODE,
                                r.NAME AS NAME,
                                r.UNIT AS UNIT,
                                NVL(r.PRICE, NVL(r.AMOUNT, 0)) AS PRICE,
                                r.TYPE_ID AS RESOURCE_TYPE_ID,
                                GREATEST(
                                    CASE WHEN UPPER(NVL(r.RESOURCE_CODE, '')) = UPPER(:q) THEN 100 ELSE 0 END,
                                    CASE WHEN UPPER(r.NAME) LIKE '%' || UPPER(:q) || '%' THEN 90 ELSE 0 END,
                                    UTL_MATCH.JARO_WINKLER_SIMILARITY(UPPER(r.NAME), UPPER(:q)),
                                    UTL_MATCH.EDIT_DISTANCE_SIMILARITY(UPPER(r.NAME), UPPER(:q))
                                ) AS SCORE
                            FROM CLS_RESOURCES r
                            WHERE r.IS_ACTIVE = 'Y'
                        )
                        WHERE SCORE >= :th
                        ORDER BY SCORE DESC
                        FETCH FIRST 1 ROWS ONLY
                        """,
                        {"q": token, "th": max(0, min(int(threshold), 100))},
                    )
                    row = _rows(r)
                    if row:
                        x = row[0]
                        matches.append(
                            {
                                "entity_type": x.get("entity_type"),
                                "id": x.get("entity_id"),
                                "code": x.get("code"),
                                "name": x.get("name"),
                                "unit": x.get("unit"),
                                "price": _safe_float(x.get("price"), 0.0),
                                "resource_type_id": x.get("resource_type_id"),
                                "qty": qty,
                                "confidence": round(_safe_float(x.get("score"), 0.0), 2),
                                "original": part["original"],
                                "token": token,
                            }
                        )
                    else:
                        matches.append(
                            {
                                "entity_type": None,
                                "id": None,
                                "code": None,
                                "name": None,
                                "unit": None,
                                "price": 0.0,
                                "qty": qty,
                                "confidence": 0,
                                "original": part["original"],
                                "token": token,
                            }
                        )

            return {"success": True, "matches": matches, "backend": "oracle"}
        except Exception as e:
            return {"success": False, "error": str(e), "matches": [], "backend": "oracle"}

    # ---------------- CRM ----------------
    @staticmethod
    def get_crm_sources() -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                data = _rows(
                    db.execute_query(
                        """
                        SELECT ID, CODE, NAME_RU, NAME_RO, SORT_ORDER, IS_ACTIVE
                        FROM CLS_CRM_SOURCES
                        WHERE IS_ACTIVE = 'Y'
                        ORDER BY SORT_ORDER, ID
                        """
                    )
                )
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_crm_stages() -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                data = _rows(
                    db.execute_query(
                        """
                        SELECT ID, CODE, NAME_RU, NAME_RO, SORT_ORDER, IS_FINAL, IS_WON, IS_LOST, IS_ACTIVE
                        FROM CLS_CRM_STAGES
                        WHERE IS_ACTIVE = 'Y'
                        ORDER BY SORT_ORDER, ID
                        """
                    )
                )
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_crm_leads(search: Optional[str] = None, stage_code: Optional[str] = None, limit: int = 300) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                sql = """
                    SELECT * FROM (
                        SELECT
                            ID, LEAD_NO, SOURCE_ID, SOURCE_CODE, SOURCE_NAME_RU, SOURCE_NAME_RO,
                            STAGE_ID, STAGE_CODE, STAGE_NAME_RU, STAGE_NAME_RO, STAGE_SORT_ORDER,
                            LANG_PREF, CONTACT_NAME, COMPANY_NAME, PHONE, EMAIL, LOCATION,
                            SUBJECT, NEEDS_TEXT, BUDGET_AMOUNT, CURRENCY, EXPECTED_CLOSE_DATE,
                            ASSIGNED_TO, PROJECT_ID, PROJECT_NAME, ESTIMATE_ID, ESTIMATE_NAME,
                            CONTRACT_ID, CONTRACT_NO, EMAIL_FROM, EMAIL_SUBJECT, EMAIL_MESSAGE_ID,
                            EMAIL_RECEIVED_AT, EXTERNAL_REF, CREATED_AT, UPDATED_AT, LAST_ACTIVITY_AT
                        FROM V_CLS_CRM_LEADS
                        WHERE IS_ACTIVE = 'Y'
                """
                params: Dict[str, Any] = {"lim": max(1, min(limit, 2000))}
                if stage_code and stage_code.strip():
                    sql += " AND UPPER(STAGE_CODE) = UPPER(:stage_code)"
                    params["stage_code"] = stage_code.strip()
                if search and search.strip():
                    sql += """
                        AND (
                            UPPER(NVL(CONTACT_NAME, '')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(COMPANY_NAME, '')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(PHONE, '')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(EMAIL, '')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(SUBJECT, '')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(LEAD_NO, '')) LIKE '%' || UPPER(:q) || '%'
                        )
                    """
                    params["q"] = search.strip()
                sql += " ORDER BY STAGE_SORT_ORDER, NVL(LAST_ACTIVITY_AT, CREATED_AT) DESC ) WHERE ROWNUM <= :lim"
                data = _rows(db.execute_query(sql, params))
                funnel = _rows(
                    db.execute_query(
                        """
                        SELECT STAGE_ID, STAGE_CODE, STAGE_NAME_RU, STAGE_NAME_RO, SORT_ORDER, LEADS_COUNT, BUDGET_TOTAL
                        FROM V_CLS_CRM_FUNNEL
                        ORDER BY SORT_ORDER
                        """
                    )
                )
                return {"success": True, "data": data, "funnel": funnel, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "funnel": [], "count": 0}

    @staticmethod
    def create_crm_lead(payload: Dict[str, Any], created_by: str = "ui") -> Dict[str, Any]:
        try:
            source_code = (payload.get("source_code") or "MANUAL").strip()
            stage_code = (payload.get("stage_code") or "NEW").strip()
            contact_name = (payload.get("contact_name") or "").strip() or None
            company_name = (payload.get("company_name") or "").strip() or None
            phone = (payload.get("phone") or "").strip() or None
            email_value = (payload.get("email") or "").strip() or None
            location = (payload.get("location") or "").strip() or None
            subject = (payload.get("subject") or "").strip() or None
            needs_text = (payload.get("needs_text") or "").strip() or None
            currency = (payload.get("currency") or "MDL").strip() or "MDL"
            lang_pref = (payload.get("lang_pref") or _guess_lang_pref((subject or "") + " " + (needs_text or ""))).strip().lower()
            budget_amount = _safe_float(payload.get("budget_amount"), 0.0)
            expected_close_date = payload.get("expected_close_date")
            assigned_to = (payload.get("assigned_to") or created_by or "ui").strip()[:120]
            project_id = payload.get("project_id")
            estimate_id = payload.get("estimate_id")
            external_ref = (payload.get("external_ref") or "").strip() or None
            email_from = (payload.get("email_from") or "").strip() or None
            email_subject = (payload.get("email_subject") or "").strip() or subject
            email_message_id = (payload.get("email_message_id") or "").strip() or None

            if not (contact_name or company_name or phone or email_value or subject):
                return {"success": False, "error": "At least one lead identity field required"}

            with DatabaseModel() as db:
                source = _rows(db.execute_query("SELECT ID FROM CLS_CRM_SOURCES WHERE UPPER(CODE)=UPPER(:c)", {"c": source_code}))
                stage = _rows(db.execute_query("SELECT ID FROM CLS_CRM_STAGES WHERE UPPER(CODE)=UPPER(:c)", {"c": stage_code}))
                if not source:
                    return {"success": False, "error": f"Unknown source_code: {source_code}"}
                if not stage:
                    return {"success": False, "error": f"Unknown stage_code: {stage_code}"}

                if email_message_id:
                    dup = _rows(
                        db.execute_query(
                            "SELECT ID, LEAD_NO FROM CLS_CRM_LEADS WHERE EMAIL_MESSAGE_ID = :mid",
                            {"mid": email_message_id},
                        )
                    )
                    if dup:
                        return {"success": True, "id": dup[0].get("id"), "lead_no": dup[0].get("lead_no"), "deduplicated": True}

                lead_no = payload.get("lead_no")
                if not lead_no:
                    lead_no = _rows(
                        db.execute_query(
                            """
                            SELECT 'LEAD-CLS-' || TO_CHAR(SYSDATE,'YYYY') || '-' ||
                                   LPAD(TO_CHAR(CLS_CRM_LEADS_SEQ.NEXTVAL), 6, '0') AS LEAD_NO
                            FROM dual
                            """
                        )
                    )[0]["lead_no"]

                db.execute_query(
                    """
                    INSERT INTO CLS_CRM_LEADS (
                        LEAD_NO, SOURCE_ID, STAGE_ID, LANG_PREF,
                        CONTACT_NAME, COMPANY_NAME, PHONE, EMAIL, LOCATION,
                        SUBJECT, NEEDS_TEXT, BUDGET_AMOUNT, CURRENCY, EXPECTED_CLOSE_DATE,
                        ASSIGNED_TO, PROJECT_ID, ESTIMATE_ID, EXTERNAL_REF,
                        EMAIL_FROM, EMAIL_SUBJECT, EMAIL_MESSAGE_ID, EMAIL_RECEIVED_AT,
                        LAST_ACTIVITY_AT, IS_ACTIVE
                    ) VALUES (
                        :lead_no, :source_id, :stage_id, :lang_pref,
                        :contact_name, :company_name, :phone, :email, :location,
                        :subject, :needs_text, :budget_amount, :currency, :expected_close_date,
                        :assigned_to, :project_id, :estimate_id, :external_ref,
                        :email_from, :email_subject, :email_message_id, :email_received_at,
                        CURRENT_TIMESTAMP, 'Y'
                    )
                    """,
                    {
                        "lead_no": lead_no,
                        "source_id": source[0]["id"],
                        "stage_id": stage[0]["id"],
                        "lang_pref": lang_pref[:2],
                        "contact_name": contact_name,
                        "company_name": company_name,
                        "phone": phone,
                        "email": email_value,
                        "location": location,
                        "subject": subject,
                        "needs_text": needs_text,
                        "budget_amount": budget_amount if budget_amount > 0 else None,
                        "currency": currency,
                        "expected_close_date": expected_close_date,
                        "assigned_to": assigned_to,
                        "project_id": project_id,
                        "estimate_id": estimate_id,
                        "external_ref": external_ref,
                        "email_from": email_from,
                        "email_subject": email_subject,
                        "email_message_id": email_message_id,
                        "email_received_at": payload.get("email_received_at"),
                    },
                )
                db.execute_query(
                    """
                    INSERT INTO CLS_CRM_ACTIVITIES (LEAD_ID, ACTIVITY_TYPE, NOTE_TEXT, PAYLOAD_JSON, CREATED_BY)
                    SELECT ID, 'LEAD_CREATED', :note, :payload, :created_by
                    FROM CLS_CRM_LEADS
                    WHERE LEAD_NO = :lead_no
                    """,
                    {
                        "lead_no": lead_no,
                        "note": "Lead captured",
                        "payload": json.dumps({"source_code": source_code, "stage_code": stage_code}, ensure_ascii=False),
                        "created_by": created_by[:120],
                    },
                )
                new_row = _rows(db.execute_query("SELECT ID FROM CLS_CRM_LEADS WHERE LEAD_NO = :x", {"x": lead_no}))
                db.connection.commit()
                return {"success": True, "id": new_row[0].get("id") if new_row else None, "lead_no": lead_no}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def update_crm_lead(lead_id: int, payload: Dict[str, Any], updated_by: str = "ui") -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                existing = _rows(
                    db.execute_query(
                        """
                        SELECT ID, STAGE_ID, STAGE_CODE, CONTRACT_ID
                        FROM V_CLS_CRM_LEADS
                        WHERE ID = :id
                        """,
                        {"id": lead_id},
                    )
                )
                if not existing:
                    return {"success": False, "error": "Lead not found"}

                current = existing[0]
                stage_id = current.get("stage_id")
                stage_code_before = current.get("stage_code")
                if payload.get("stage_code"):
                    stage = _rows(
                        db.execute_query(
                            "SELECT ID, CODE FROM CLS_CRM_STAGES WHERE UPPER(CODE)=UPPER(:c)",
                            {"c": str(payload.get("stage_code"))},
                        )
                    )
                    if not stage:
                        return {"success": False, "error": "Unknown stage_code"}
                    stage_id = stage[0]["id"]

                db.execute_query(
                    """
                    UPDATE CLS_CRM_LEADS
                    SET STAGE_ID = :stage_id,
                        CONTACT_NAME = NVL(:contact_name, CONTACT_NAME),
                        COMPANY_NAME = NVL(:company_name, COMPANY_NAME),
                        PHONE = NVL(:phone, PHONE),
                        EMAIL = NVL(:email, EMAIL),
                        LOCATION = NVL(:location, LOCATION),
                        SUBJECT = NVL(:subject, SUBJECT),
                        NEEDS_TEXT = NVL(:needs_text, NEEDS_TEXT),
                        BUDGET_AMOUNT = :budget_amount,
                        CURRENCY = NVL(:currency, CURRENCY),
                        EXPECTED_CLOSE_DATE = :expected_close_date,
                        ASSIGNED_TO = NVL(:assigned_to, ASSIGNED_TO),
                        PROJECT_ID = :project_id,
                        ESTIMATE_ID = :estimate_id,
                        LAST_ACTIVITY_AT = CURRENT_TIMESTAMP,
                        UPDATED_AT = CURRENT_TIMESTAMP
                    WHERE ID = :id
                    """,
                    {
                        "id": lead_id,
                        "stage_id": stage_id,
                        "contact_name": (payload.get("contact_name") or "").strip() or None,
                        "company_name": (payload.get("company_name") or "").strip() or None,
                        "phone": (payload.get("phone") or "").strip() or None,
                        "email": (payload.get("email") or "").strip() or None,
                        "location": (payload.get("location") or "").strip() or None,
                        "subject": (payload.get("subject") or "").strip() or None,
                        "needs_text": (payload.get("needs_text") or "").strip() or None,
                        "budget_amount": _safe_float(payload.get("budget_amount"), 0.0) if payload.get("budget_amount") not in (None, "") else None,
                        "currency": (payload.get("currency") or "").strip() or None,
                        "expected_close_date": payload.get("expected_close_date"),
                        "assigned_to": (payload.get("assigned_to") or "").strip() or None,
                        "project_id": payload.get("project_id"),
                        "estimate_id": payload.get("estimate_id"),
                    },
                )
                note = (payload.get("activity_note") or "").strip()
                if payload.get("stage_code") and str(payload.get("stage_code")).upper() != str(stage_code_before or "").upper():
                    note = note or f"Stage changed: {stage_code_before} -> {payload.get('stage_code')}"
                if note:
                    db.execute_query(
                        """
                        INSERT INTO CLS_CRM_ACTIVITIES (LEAD_ID, ACTIVITY_TYPE, NOTE_TEXT, PAYLOAD_JSON, CREATED_BY)
                        VALUES (:lead_id, 'UPDATE', :note_text, :payload_json, :created_by)
                        """,
                        {
                            "lead_id": lead_id,
                            "note_text": note,
                            "payload_json": json.dumps(payload, ensure_ascii=False),
                            "created_by": updated_by[:120],
                        },
                    )
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_crm_activities(lead_id: int, limit: int = 100) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                data = _rows(
                    db.execute_query(
                        """
                        SELECT * FROM (
                            SELECT ID, LEAD_ID, ACTIVITY_TYPE, NOTE_TEXT, PAYLOAD_JSON, CREATED_BY, CREATED_AT
                            FROM CLS_CRM_ACTIVITIES
                            WHERE LEAD_ID = :lead_id
                            ORDER BY CREATED_AT DESC, ID DESC
                        ) WHERE ROWNUM <= :lim
                        """,
                        {"lead_id": lead_id, "lim": max(1, min(limit, 500))},
                    )
                )
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def add_crm_activity(lead_id: int, payload: Dict[str, Any], created_by: str = "ui") -> Dict[str, Any]:
        try:
            activity_type = (payload.get("activity_type") or "NOTE").strip()[:60]
            note_text = (payload.get("note_text") or payload.get("note") or "").strip()
            if not note_text:
                return {"success": False, "error": "note_text required"}
            with DatabaseModel() as db:
                db.execute_query(
                    """
                    INSERT INTO CLS_CRM_ACTIVITIES (LEAD_ID, ACTIVITY_TYPE, NOTE_TEXT, PAYLOAD_JSON, CREATED_BY)
                    VALUES (:lead_id, :activity_type, :note_text, :payload_json, :created_by)
                    """,
                    {
                        "lead_id": lead_id,
                        "activity_type": activity_type,
                        "note_text": note_text,
                        "payload_json": json.dumps(payload.get("payload") or {}, ensure_ascii=False),
                        "created_by": created_by[:120],
                    },
                )
                db.execute_query(
                    "UPDATE CLS_CRM_LEADS SET LAST_ACTIVITY_AT = CURRENT_TIMESTAMP, UPDATED_AT = CURRENT_TIMESTAMP WHERE ID = :id",
                    {"id": lead_id},
                )
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def register_crm_contract(lead_id: int, payload: Dict[str, Any], created_by: str = "ui") -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                lead_rows = _rows(
                    db.execute_query(
                        """
                        SELECT ID, LEAD_NO, PROJECT_ID, ESTIMATE_ID, BUDGET_AMOUNT, CURRENCY, CONTRACT_ID
                        FROM CLS_CRM_LEADS
                        WHERE ID = :id
                        """,
                        {"id": lead_id},
                    )
                )
                if not lead_rows:
                    return {"success": False, "error": "Lead not found"}
                lead = lead_rows[0]
                if lead.get("contract_id"):
                    return {"success": False, "error": "Contract already registered for this lead"}

                contract_no = (payload.get("contract_no") or "").strip()
                if not contract_no:
                    contract_no = _rows(
                        db.execute_query(
                            """
                            SELECT 'CON-CLS-' || TO_CHAR(SYSDATE,'YYYY') || '-' ||
                                   LPAD(TO_CHAR(CLS_CRM_CONTRACTS_SEQ.NEXTVAL), 6, '0') AS CONTRACT_NO
                            FROM dual
                            """
                        )
                    )[0]["contract_no"]

                amount = _safe_float(payload.get("amount"), _safe_float(lead.get("budget_amount"), 0.0))
                currency = (payload.get("currency") or lead.get("currency") or "MDL").strip()
                registry_no = (payload.get("registry_no") or "").strip() or None
                status = (payload.get("status") or "REGISTERED").strip()
                notes = (payload.get("notes") or "").strip() or None

                db.execute_query(
                    """
                    INSERT INTO CLS_CRM_CONTRACTS (
                        CONTRACT_NO, LEAD_ID, PROJECT_ID, ESTIMATE_ID,
                        CONTRACT_DATE, AMOUNT, CURRENCY, STATUS, REGISTRY_NO, NOTES, CREATED_BY
                    ) VALUES (
                        :contract_no, :lead_id, :project_id, :estimate_id,
                        :contract_date, :amount, :currency, :status, :registry_no, :notes, :created_by
                    )
                    """,
                    {
                        "contract_no": contract_no,
                        "lead_id": lead_id,
                        "project_id": payload.get("project_id") or lead.get("project_id"),
                        "estimate_id": payload.get("estimate_id") or lead.get("estimate_id"),
                        "contract_date": payload.get("contract_date"),
                        "amount": amount,
                        "currency": currency,
                        "status": status,
                        "registry_no": registry_no,
                        "notes": notes,
                        "created_by": created_by[:120],
                    },
                )
                contract_row = _rows(
                    db.execute_query(
                        "SELECT ID FROM CLS_CRM_CONTRACTS WHERE CONTRACT_NO = :x",
                        {"x": contract_no},
                    )
                )
                contract_id = contract_row[0]["id"] if contract_row else None
                won_stage = _rows(
                    db.execute_query(
                        "SELECT ID FROM CLS_CRM_STAGES WHERE CODE = 'CONTRACT_REGISTERED'",
                    )
                )
                won_stage_id = won_stage[0]["id"] if won_stage else lead.get("stage_id")
                db.execute_query(
                    """
                    UPDATE CLS_CRM_LEADS
                    SET CONTRACT_ID = :contract_id,
                        STAGE_ID = :stage_id,
                        LAST_ACTIVITY_AT = CURRENT_TIMESTAMP,
                        UPDATED_AT = CURRENT_TIMESTAMP
                    WHERE ID = :lead_id
                    """,
                    {"contract_id": contract_id, "stage_id": won_stage_id, "lead_id": lead_id},
                )
                db.execute_query(
                    """
                    INSERT INTO CLS_CRM_ACTIVITIES (LEAD_ID, ACTIVITY_TYPE, NOTE_TEXT, PAYLOAD_JSON, CREATED_BY)
                    VALUES (:lead_id, 'CONTRACT_REGISTERED', :note_text, :payload_json, :created_by)
                    """,
                    {
                        "lead_id": lead_id,
                        "note_text": f"Contract registered: {contract_no}",
                        "payload_json": json.dumps({"contract_id": contract_id, "contract_no": contract_no}, ensure_ascii=False),
                        "created_by": created_by[:120],
                    },
                )

                # Optional: auto-create shell-native Colass contract record linked to CRM + estimate.
                create_shell_contract = payload.get("create_shell_contract", True)
                shell_contract_id = None
                if create_shell_contract:
                    try:
                        shell_out = ColassController.create_contract(
                            {
                                "contract_no": contract_no,
                                "contract_date": payload.get("contract_date"),
                                "currency": currency,
                                "total_amount": amount,
                                "status": "REGISTERED",
                                "project_id": payload.get("project_id") or lead.get("project_id"),
                                "estimate_id": payload.get("estimate_id") or lead.get("estimate_id"),
                                "crm_lead_id": lead_id,
                                "crm_contract_id": contract_id,
                                "company_name": payload.get("company_name"),
                                "notes": notes,
                            },
                            created_by=created_by,
                            db=db,
                            commit=False,
                        )
                        if shell_out.get("success"):
                            shell_contract_id = shell_out.get("id")
                    except Exception:
                        shell_contract_id = None
                db.connection.commit()
                return {
                    "success": True,
                    "contract_id": contract_id,
                    "contract_no": contract_no,
                    "shell_contract_id": shell_contract_id,
                }
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_crm_contracts(limit: int = 200) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                data = _rows(
                    db.execute_query(
                        """
                        SELECT * FROM (
                            SELECT
                                c.ID, c.CONTRACT_NO, c.LEAD_ID, l.LEAD_NO,
                                l.CONTACT_NAME, l.COMPANY_NAME, c.PROJECT_ID, c.ESTIMATE_ID,
                                c.CONTRACT_DATE, c.AMOUNT, c.CURRENCY, c.STATUS, c.REGISTRY_NO,
                                c.REGISTERED_AT, c.CREATED_AT
                            FROM CLS_CRM_CONTRACTS c
                            JOIN CLS_CRM_LEADS l ON l.ID = c.LEAD_ID
                            ORDER BY c.CONTRACT_DATE DESC, c.ID DESC
                        ) WHERE ROWNUM <= :lim
                        """,
                        {"lim": max(1, min(limit, 1000))},
                    )
                )
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    # ---------------- Colass Contracts (shell-native) ----------------
    @staticmethod
    def _to_int(value: Any) -> Optional[int]:
        try:
            if value is None or value == "":
                return None
            return int(value)
        except Exception:
            return None

    @staticmethod
    def _sync_contract_items_from_estimate(db: DatabaseModel, contract_id: int, estimate_id: Optional[int]) -> int:
        if not estimate_id:
            return 0
        db.execute_query("DELETE FROM CLS_CONTRACT_ITEMS WHERE CONTRACT_ID = :id", {"id": contract_id})
        r = db.execute_query(
            """
            INSERT INTO CLS_CONTRACT_ITEMS (
                CONTRACT_ID, ESTIMATE_ITEM_ID, RESOURCE_TYPE_ID, ITEM_CODE, ITEM_NAME, UNIT,
                QTY, PRICE, AMOUNT, SOURCE_DOC, SOURCE_ROW
            )
            SELECT
                :contract_id, i.ID, i.RESOURCE_TYPE_ID, i.ITEM_CODE, i.ITEM_NAME, i.UNIT,
                i.QTY, i.PRICE, i.AMOUNT, i.SOURCE_DOC, i.SOURCE_ROW
            FROM CLS_ESTIMATE_ITEMS i
            WHERE i.ESTIMATE_ID = :estimate_id
            """,
            {"contract_id": contract_id, "estimate_id": estimate_id},
        )
        return int(r.get("rowcount") or 0)

    @staticmethod
    def get_contracts(search: Optional[str] = None, limit: int = 300) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                sql = """
                    SELECT * FROM (
                        SELECT
                            c.ID, c.CONTRACT_NO, c.CONTRACT_DATE, c.LEGAL_FORM, c.COMPANY_NAME, c.ADDRESS,
                            c.FISCAL_CODE, c.VAT_CODE, c.BANK_NAME, c.BANK_BRANCH, c.BANK_BIC,
                            c.IBAN_MDL, c.IBAN_EUR, c.IBAN_USD,
                            c.CURRENCY, c.TOTAL_AMOUNT, c.STATUS, c.NOTES,
                            c.PROJECT_ID, c.PROJECT_NAME, c.ESTIMATE_ID, c.ESTIMATE_NAME,
                            c.CRM_LEAD_ID, c.LEAD_NO, c.LEAD_CONTACT_NAME, c.LEAD_COMPANY_NAME,
                            c.CRM_CONTRACT_ID, c.CRM_CONTRACT_NO,
                            c.ITEMS_COUNT, c.ITEMS_TOTAL_AMOUNT,
                            c.CREATED_AT, c.UPDATED_AT
                        FROM V_CLS_CONTRACTS c
                        WHERE 1=1
                """
                params: Dict[str, Any] = {"lim": max(1, min(limit, 2000))}
                if search and search.strip():
                    sql += """
                        AND (
                            UPPER(NVL(c.CONTRACT_NO,'')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(c.COMPANY_NAME,'')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(c.LEAD_CONTACT_NAME,'')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(c.LEAD_COMPANY_NAME,'')) LIKE '%' || UPPER(:q) || '%'
                            OR UPPER(NVL(c.FISCAL_CODE,'')) LIKE '%' || UPPER(:q) || '%'
                        )
                    """
                    params["q"] = search.strip()
                sql += " ORDER BY c.CONTRACT_DATE DESC, c.ID DESC ) WHERE ROWNUM <= :lim"
                data = _rows(db.execute_query(sql, params))
                return {"success": True, "data": data, "count": len(data)}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "count": 0}

    @staticmethod
    def get_contract_detail(contract_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                rows = _rows(
                    db.execute_query(
                        "SELECT * FROM V_CLS_CONTRACTS WHERE ID = :id",
                        {"id": contract_id},
                    )
                )
                if not rows:
                    return {"success": False, "error": "Contract not found"}
                emails = _rows(
                    db.execute_query(
                        "SELECT ID, CONTRACT_ID, SORT_ORDER, EMAIL, IS_PRIMARY FROM CLS_CONTRACT_EMAILS WHERE CONTRACT_ID = :id ORDER BY SORT_ORDER, ID",
                        {"id": contract_id},
                    )
                )
                phones = _rows(
                    db.execute_query(
                        "SELECT ID, CONTRACT_ID, SORT_ORDER, PHONE, IS_PRIMARY FROM CLS_CONTRACT_PHONES WHERE CONTRACT_ID = :id ORDER BY SORT_ORDER, ID",
                        {"id": contract_id},
                    )
                )
                routes = _rows(
                    db.execute_query(
                        "SELECT ID, CONTRACT_ID, SORT_ORDER, ROUTE_NAME FROM CLS_CONTRACT_ROUTES WHERE CONTRACT_ID = :id ORDER BY SORT_ORDER, ID",
                        {"id": contract_id},
                    )
                )
                items = _rows(
                    db.execute_query(
                        """
                        SELECT i.ID, i.CONTRACT_ID, i.ESTIMATE_ITEM_ID, i.RESOURCE_TYPE_ID, rt.CODE AS TYPE_CODE,
                               rt.NAME_RU AS TYPE_NAME_RU, rt.NAME_RO AS TYPE_NAME_RO,
                               i.ITEM_CODE, i.ITEM_NAME, i.UNIT, i.QTY, i.PRICE, i.AMOUNT, i.SOURCE_DOC, i.SOURCE_ROW
                        FROM CLS_CONTRACT_ITEMS i
                        LEFT JOIN CLS_RESOURCE_TYPES rt ON rt.ID = i.RESOURCE_TYPE_ID
                        WHERE i.CONTRACT_ID = :id
                        ORDER BY i.ID
                        """,
                        {"id": contract_id},
                    )
                )
                return {"success": True, "data": rows[0], "emails": emails, "phones": phones, "routes": routes, "items": items}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def create_contract(payload: Dict[str, Any], created_by: str = "ui", db: Optional[DatabaseModel] = None, commit: bool = True) -> Dict[str, Any]:
        own_db = db is None
        ctx = None
        try:
            ctx = DatabaseModel() if own_db else None
            _db = ctx.__enter__() if own_db and ctx else db
            assert _db is not None

            contract_no = (payload.get("contract_no") or "").strip()
            if not contract_no:
                seq = _rows(
                    _db.execute_query(
                        """
                        SELECT 'CLS-CON-' || TO_CHAR(SYSDATE,'YYYY') || '-' ||
                               LPAD(TO_CHAR(CLS_CONTRACTS_SEQ.NEXTVAL), 6, '0') AS CONTRACT_NO
                        FROM dual
                        """
                    )
                )
                contract_no = seq[0]["contract_no"]

            estimate_id = ColassController._to_int(payload.get("estimate_id"))
            total_amount = _safe_float(payload.get("total_amount"), 0.0)
            if total_amount <= 0 and estimate_id:
                est = _rows(_db.execute_query("SELECT TOTAL_AMOUNT FROM CLS_ESTIMATES WHERE ID = :id", {"id": estimate_id}))
                total_amount = _safe_float(est[0]["total_amount"], 0.0) if est else 0.0

            _db.execute_query(
                """
                INSERT INTO CLS_CONTRACTS (
                    CONTRACT_NO, CONTRACT_DATE, LEGAL_FORM, COMPANY_NAME, ADDRESS,
                    FISCAL_CODE, VAT_CODE, BANK_NAME, BANK_BRANCH, BANK_BIC,
                    IBAN_MDL, IBAN_EUR, IBAN_USD,
                    CURRENCY, TOTAL_AMOUNT, STATUS, NOTES,
                    PROJECT_ID, ESTIMATE_ID, CRM_LEAD_ID, CRM_CONTRACT_ID, CREATED_BY
                ) VALUES (
                    :contract_no, :contract_date, :legal_form, :company_name, :address,
                    :fiscal_code, :vat_code, :bank_name, :bank_branch, :bank_bic,
                    :iban_mdl, :iban_eur, :iban_usd,
                    :currency, :total_amount, :status, :notes,
                    :project_id, :estimate_id, :crm_lead_id, :crm_contract_id, :created_by
                )
                """,
                {
                    "contract_no": contract_no,
                    "contract_date": payload.get("contract_date"),
                    "legal_form": (payload.get("legal_form") or "").strip() or None,
                    "company_name": (payload.get("company_name") or "").strip() or None,
                    "address": (payload.get("address") or "").strip() or None,
                    "fiscal_code": (payload.get("fiscal_code") or "").strip() or None,
                    "vat_code": (payload.get("vat_code") or "").strip() or None,
                    "bank_name": (payload.get("bank_name") or "").strip() or None,
                    "bank_branch": (payload.get("bank_branch") or "").strip() or None,
                    "bank_bic": (payload.get("bank_bic") or "").strip() or None,
                    "iban_mdl": (payload.get("iban_mdl") or "").strip() or None,
                    "iban_eur": (payload.get("iban_eur") or "").strip() or None,
                    "iban_usd": (payload.get("iban_usd") or "").strip() or None,
                    "currency": (payload.get("currency") or "MDL").strip() or "MDL",
                    "total_amount": total_amount,
                    "status": (payload.get("status") or "DRAFT").strip() or "DRAFT",
                    "notes": (payload.get("notes") or "").strip() or None,
                    "project_id": ColassController._to_int(payload.get("project_id")),
                    "estimate_id": estimate_id,
                    "crm_lead_id": ColassController._to_int(payload.get("crm_lead_id")),
                    "crm_contract_id": ColassController._to_int(payload.get("crm_contract_id")),
                    "created_by": (created_by or "ui")[:120],
                },
            )
            row = _rows(_db.execute_query("SELECT ID FROM CLS_CONTRACTS WHERE CONTRACT_NO = :x", {"x": contract_no}))
            contract_id = row[0]["id"] if row else None

            inserted_items = 0
            if contract_id and estimate_id:
                inserted_items = ColassController._sync_contract_items_from_estimate(_db, int(contract_id), estimate_id)

            if commit:
                _db.connection.commit()
            if own_db and ctx:
                ctx.__exit__(None, None, None)
            return {"success": True, "id": contract_id, "contract_no": contract_no, "inserted_items": inserted_items}
        except Exception as e:
            try:
                if own_db and ctx:
                    ctx.__exit__(None, None, None)
            except Exception:
                pass
            return {"success": False, "error": str(e)}

    @staticmethod
    def update_contract(contract_id: int, payload: Dict[str, Any]) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                estimate_id = ColassController._to_int(payload.get("estimate_id"))
                db.execute_query(
                    """
                    UPDATE CLS_CONTRACTS
                    SET CONTRACT_DATE = :contract_date,
                        LEGAL_FORM = :legal_form,
                        COMPANY_NAME = :company_name,
                        ADDRESS = :address,
                        FISCAL_CODE = :fiscal_code,
                        VAT_CODE = :vat_code,
                        BANK_NAME = :bank_name,
                        BANK_BRANCH = :bank_branch,
                        BANK_BIC = :bank_bic,
                        IBAN_MDL = :iban_mdl,
                        IBAN_EUR = :iban_eur,
                        IBAN_USD = :iban_usd,
                        CURRENCY = :currency,
                        TOTAL_AMOUNT = :total_amount,
                        STATUS = :status,
                        NOTES = :notes,
                        PROJECT_ID = :project_id,
                        ESTIMATE_ID = :estimate_id,
                        CRM_LEAD_ID = :crm_lead_id,
                        UPDATED_AT = CURRENT_TIMESTAMP
                    WHERE ID = :id
                    """,
                    {
                        "id": contract_id,
                        "contract_date": payload.get("contract_date"),
                        "legal_form": (payload.get("legal_form") or "").strip() or None,
                        "company_name": (payload.get("company_name") or "").strip() or None,
                        "address": (payload.get("address") or "").strip() or None,
                        "fiscal_code": (payload.get("fiscal_code") or "").strip() or None,
                        "vat_code": (payload.get("vat_code") or "").strip() or None,
                        "bank_name": (payload.get("bank_name") or "").strip() or None,
                        "bank_branch": (payload.get("bank_branch") or "").strip() or None,
                        "bank_bic": (payload.get("bank_bic") or "").strip() or None,
                        "iban_mdl": (payload.get("iban_mdl") or "").strip() or None,
                        "iban_eur": (payload.get("iban_eur") or "").strip() or None,
                        "iban_usd": (payload.get("iban_usd") or "").strip() or None,
                        "currency": (payload.get("currency") or "MDL").strip(),
                        "total_amount": _safe_float(payload.get("total_amount"), 0.0),
                        "status": (payload.get("status") or "DRAFT").strip(),
                        "notes": (payload.get("notes") or "").strip() or None,
                        "project_id": ColassController._to_int(payload.get("project_id")),
                        "estimate_id": estimate_id,
                        "crm_lead_id": ColassController._to_int(payload.get("crm_lead_id")),
                    },
                )
                if payload.get("sync_items_from_estimate") and estimate_id:
                    ColassController._sync_contract_items_from_estimate(db, contract_id, estimate_id)
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def delete_contract(contract_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                db.execute_query("DELETE FROM CLS_CONTRACT_EMAILS WHERE CONTRACT_ID = :id", {"id": contract_id})
                db.execute_query("DELETE FROM CLS_CONTRACT_PHONES WHERE CONTRACT_ID = :id", {"id": contract_id})
                db.execute_query("DELETE FROM CLS_CONTRACT_ROUTES WHERE CONTRACT_ID = :id", {"id": contract_id})
                db.execute_query("DELETE FROM CLS_CONTRACT_ITEMS WHERE CONTRACT_ID = :id", {"id": contract_id})
                db.execute_query("DELETE FROM CLS_CONTRACTS WHERE ID = :id", {"id": contract_id})
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def add_contract_contact(contract_id: int, kind: str, value: str, is_primary: str = "N") -> Dict[str, Any]:
        try:
            v = (value or "").strip()
            if not v:
                return {"success": False, "error": "value required"}
            k = (kind or "").strip().lower()
            if k not in ("email", "phone", "route"):
                return {"success": False, "error": "kind must be email|phone|route"}
            with DatabaseModel() as db:
                if k == "email":
                    sort_row = _rows(db.execute_query("SELECT NVL(MAX(SORT_ORDER),0)+1 AS N FROM CLS_CONTRACT_EMAILS WHERE CONTRACT_ID=:id", {"id": contract_id}))
                    db.execute_query(
                        "INSERT INTO CLS_CONTRACT_EMAILS (CONTRACT_ID, SORT_ORDER, EMAIL, IS_PRIMARY) VALUES (:id,:s,:v,:p)",
                        {"id": contract_id, "s": sort_row[0]["n"] if sort_row else 1, "v": v, "p": "Y" if is_primary == "Y" else "N"},
                    )
                elif k == "phone":
                    sort_row = _rows(db.execute_query("SELECT NVL(MAX(SORT_ORDER),0)+1 AS N FROM CLS_CONTRACT_PHONES WHERE CONTRACT_ID=:id", {"id": contract_id}))
                    db.execute_query(
                        "INSERT INTO CLS_CONTRACT_PHONES (CONTRACT_ID, SORT_ORDER, PHONE, IS_PRIMARY) VALUES (:id,:s,:v,:p)",
                        {"id": contract_id, "s": sort_row[0]["n"] if sort_row else 1, "v": v, "p": "Y" if is_primary == "Y" else "N"},
                    )
                else:
                    sort_row = _rows(db.execute_query("SELECT NVL(MAX(SORT_ORDER),0)+1 AS N FROM CLS_CONTRACT_ROUTES WHERE CONTRACT_ID=:id", {"id": contract_id}))
                    db.execute_query(
                        "INSERT INTO CLS_CONTRACT_ROUTES (CONTRACT_ID, SORT_ORDER, ROUTE_NAME) VALUES (:id,:s,:v)",
                        {"id": contract_id, "s": sort_row[0]["n"] if sort_row else 1, "v": v},
                    )
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def delete_contract_contact(contract_id: int, kind: str, row_id: int) -> Dict[str, Any]:
        try:
            k = (kind or "").strip().lower()
            table = {
                "email": "CLS_CONTRACT_EMAILS",
                "phone": "CLS_CONTRACT_PHONES",
                "route": "CLS_CONTRACT_ROUTES",
            }.get(k)
            if not table:
                return {"success": False, "error": "kind must be email|phone|route"}
            with DatabaseModel() as db:
                db.execute_query(f"DELETE FROM {table} WHERE CONTRACT_ID = :cid AND ID = :id", {"cid": contract_id, "id": row_id})
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_contract_attachments(contract_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                rows = _rows(
                    db.execute_query(
                        """
                        SELECT ID, CONTRACT_ID, TYPE_CODE, TYPE_NAME_RU, TYPE_NAME_RO,
                               FILE_NAME, MIME_TYPE, FILE_SIZE, IS_ACTIVE, UPLOADED_BY, CREATED_AT
                        FROM CLS_CONTRACT_ATTACHMENTS
                        WHERE CONTRACT_ID = :id AND IS_ACTIVE = 'Y'
                        ORDER BY ID
                        """,
                        {"id": contract_id},
                    )
                )
                check = _rows(
                    db.execute_query(
                        """
                        SELECT CONTRACT_ID, HAS_FINANCIAL_TERMS, HAS_ESTIMATE_TERMS, HAS_PRICE_LIST, ATTACHMENTS_COUNT
                        FROM V_CLS_CONTRACT_ATTACHMENT_CHECK
                        WHERE CONTRACT_ID = :id
                        """,
                        {"id": contract_id},
                    )
                )
                required = check[0] if check else {}
                return {"success": True, "data": rows, "required": required}
        except Exception as e:
            return {"success": False, "error": str(e), "data": [], "required": {}}

    @staticmethod
    def add_contract_attachment(
        contract_id: int,
        type_code: str,
        file_name: str,
        mime_type: str,
        file_bytes: bytes,
        uploaded_by: str = "ui",
    ) -> Dict[str, Any]:
        try:
            t = (type_code or "").strip().upper()
            if t not in ("FINANCIAL_TERMS", "ESTIMATE_TERMS", "PRICE_LIST", "OTHER"):
                return {"success": False, "error": "Invalid type_code"}
            if not file_bytes:
                return {"success": False, "error": "Empty file"}
            names = {
                "FINANCIAL_TERMS": ("Финансовые условия", "Conditii financiare"),
                "ESTIMATE_TERMS": ("Сметные условия", "Conditii de deviz"),
                "PRICE_LIST": ("Общий прайс-лист", "Pricelist general"),
                "OTHER": ("Приложение", "Anexa"),
            }
            with DatabaseModel() as db:
                db.execute_query(
                    """
                    INSERT INTO CLS_CONTRACT_ATTACHMENTS (
                        CONTRACT_ID, TYPE_CODE, TYPE_NAME_RU, TYPE_NAME_RO,
                        FILE_NAME, MIME_TYPE, FILE_SIZE, FILE_BLOB, IS_ACTIVE, UPLOADED_BY
                    ) VALUES (
                        :contract_id, :type_code, :ru, :ro,
                        :file_name, :mime_type, :file_size, :file_blob, 'Y', :uploaded_by
                    )
                    """,
                    {
                        "contract_id": contract_id,
                        "type_code": t,
                        "ru": names[t][0],
                        "ro": names[t][1],
                        "file_name": (file_name or "attachment.bin")[:500],
                        "mime_type": (mime_type or "application/octet-stream")[:200],
                        "file_size": len(file_bytes),
                        "file_blob": file_bytes,
                        "uploaded_by": (uploaded_by or "ui")[:120],
                    },
                )
                rid = _rows(
                    db.execute_query(
                        "SELECT MAX(ID) AS ID FROM CLS_CONTRACT_ATTACHMENTS WHERE CONTRACT_ID = :id",
                        {"id": contract_id},
                    )
                )
                db.connection.commit()
                return {"success": True, "id": rid[0].get("id") if rid else None}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_contract_attachment_blob(attachment_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                with db.connection.cursor() as cur:
                    cur.execute(
                        """
                        SELECT ID, CONTRACT_ID, TYPE_CODE, FILE_NAME, MIME_TYPE, FILE_BLOB
                        FROM CLS_CONTRACT_ATTACHMENTS
                        WHERE ID = :id AND IS_ACTIVE = 'Y'
                        """,
                        {"id": attachment_id},
                    )
                    row = cur.fetchone()
                    if not row:
                        return {"success": False, "error": "Attachment not found"}
                    blob = row[5]
                    content: bytes
                    if hasattr(blob, "read"):
                        raw = blob.read()
                        content = raw if isinstance(raw, (bytes, bytearray)) else str(raw).encode("utf-8", errors="ignore")
                    elif isinstance(blob, (bytes, bytearray)):
                        content = bytes(blob)
                    else:
                        content = str(blob).encode("utf-8", errors="ignore")
                    return {
                        "success": True,
                        "id": row[0],
                        "contract_id": row[1],
                        "type_code": row[2],
                        "file_name": row[3],
                        "mime_type": row[4] or "application/octet-stream",
                        "content": content,
                    }
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def delete_contract_attachment(contract_id: int, attachment_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                db.execute_query(
                    """
                    UPDATE CLS_CONTRACT_ATTACHMENTS
                    SET IS_ACTIVE = 'N', UPDATED_AT = CURRENT_TIMESTAMP
                    WHERE CONTRACT_ID = :contract_id AND ID = :id
                    """,
                    {"contract_id": contract_id, "id": attachment_id},
                )
                db.connection.commit()
                return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def start_contract_approval(
        contract_id: int,
        approvers: List[Dict[str, Any]],
        started_by: str = "ui",
        comment_text: Optional[str] = None,
    ) -> Dict[str, Any]:
        try:
            if not approvers:
                return {"success": False, "error": "Approvers list required"}
            with DatabaseModel() as db:
                req = _rows(
                    db.execute_query(
                        """
                        SELECT HAS_FINANCIAL_TERMS, HAS_ESTIMATE_TERMS, HAS_PRICE_LIST
                        FROM V_CLS_CONTRACT_ATTACHMENT_CHECK
                        WHERE CONTRACT_ID = :id
                        """,
                        {"id": contract_id},
                    )
                )
                if not req:
                    return {"success": False, "error": "Contract not found"}
                r = req[0]
                if int(r.get("has_financial_terms") or 0) == 0:
                    return {"success": False, "error": "Missing required attachment: FINANCIAL_TERMS"}
                if int(r.get("has_estimate_terms") or 0) == 0:
                    return {"success": False, "error": "Missing required attachment: ESTIMATE_TERMS"}
                if int(r.get("has_price_list") or 0) == 0:
                    return {"success": False, "error": "Missing required attachment: PRICE_LIST"}

                active = _rows(
                    db.execute_query(
                        """
                        SELECT ID FROM CLS_CONTRACT_APPROVALS
                        WHERE CONTRACT_ID = :id AND STATUS = 'PENDING'
                        ORDER BY ID DESC
                        FETCH FIRST 1 ROWS ONLY
                        """,
                        {"id": contract_id},
                    )
                )
                if active:
                    return {"success": False, "error": "Pending approval already exists", "approval_id": active[0].get("id")}

                db.execute_query(
                    """
                    INSERT INTO CLS_CONTRACT_APPROVALS (CONTRACT_ID, STATUS, STARTED_BY, COMMENT_TEXT)
                    VALUES (:contract_id, 'PENDING', :started_by, :comment_text)
                    """,
                    {"contract_id": contract_id, "started_by": started_by[:120], "comment_text": (comment_text or "")[:2000] or None},
                )
                aid = _rows(
                    db.execute_query(
                        """
                        SELECT ID FROM CLS_CONTRACT_APPROVALS
                        WHERE CONTRACT_ID = :id
                        ORDER BY ID DESC FETCH FIRST 1 ROWS ONLY
                        """,
                        {"id": contract_id},
                    )
                )
                approval_id = int(aid[0]["id"]) if aid else 0

                for i, a in enumerate(approvers, start=1):
                    nm = (a.get("name") or "").strip()
                    if not nm:
                        continue
                    db.execute_query(
                        """
                        INSERT INTO CLS_CONTRACT_APPROVAL_STEPS (APPROVAL_ID, STEP_NO, APPROVER_NAME, APPROVER_ROLE, DECISION)
                        VALUES (:approval_id, :step_no, :approver_name, :approver_role, 'PENDING')
                        """,
                        {
                            "approval_id": approval_id,
                            "step_no": i,
                            "approver_name": nm[:300],
                            "approver_role": ((a.get("role") or "").strip() or None),
                        },
                    )
                db.execute_query(
                    "UPDATE CLS_CONTRACTS SET STATUS = 'UNDER_APPROVAL', UPDATED_AT = CURRENT_TIMESTAMP WHERE ID = :id",
                    {"id": contract_id},
                )
                db.connection.commit()
                return {"success": True, "approval_id": approval_id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_contract_approval(contract_id: int) -> Dict[str, Any]:
        try:
            with DatabaseModel() as db:
                app = _rows(
                    db.execute_query(
                        """
                        SELECT ID, CONTRACT_ID, STATUS, STARTED_BY, STARTED_AT, FINISHED_AT, COMMENT_TEXT
                        FROM CLS_CONTRACT_APPROVALS
                        WHERE CONTRACT_ID = :id
                        ORDER BY ID DESC
                        FETCH FIRST 1 ROWS ONLY
                        """,
                        {"id": contract_id},
                    )
                )
                if not app:
                    return {"success": True, "approval": None, "steps": []}
                approval = app[0]
                steps = _rows(
                    db.execute_query(
                        """
                        SELECT ID, APPROVAL_ID, STEP_NO, APPROVER_NAME, APPROVER_ROLE, DECISION, DECIDED_BY, DECIDED_AT, COMMENT_TEXT
                        FROM CLS_CONTRACT_APPROVAL_STEPS
                        WHERE APPROVAL_ID = :id
                        ORDER BY STEP_NO, ID
                        """,
                        {"id": approval["id"]},
                    )
                )
                return {"success": True, "approval": approval, "steps": steps}
        except Exception as e:
            return {"success": False, "error": str(e), "approval": None, "steps": []}

    @staticmethod
    def decide_contract_approval_step(step_id: int, decision: str, decided_by: str = "ui", comment_text: Optional[str] = None) -> Dict[str, Any]:
        try:
            d = (decision or "").strip().upper()
            if d not in ("APPROVED", "REJECTED"):
                return {"success": False, "error": "decision must be APPROVED or REJECTED"}
            with DatabaseModel() as db:
                step = _rows(
                    db.execute_query(
                        """
                        SELECT ID, APPROVAL_ID, DECISION
                        FROM CLS_CONTRACT_APPROVAL_STEPS
                        WHERE ID = :id
                        """,
                        {"id": step_id},
                    )
                )
                if not step:
                    return {"success": False, "error": "Step not found"}
                if str(step[0].get("decision") or "").upper() != "PENDING":
                    return {"success": False, "error": "Step already decided"}
                approval_id = int(step[0]["approval_id"])

                db.execute_query(
                    """
                    UPDATE CLS_CONTRACT_APPROVAL_STEPS
                    SET DECISION = :decision,
                        DECIDED_BY = :decided_by,
                        DECIDED_AT = CURRENT_TIMESTAMP,
                        COMMENT_TEXT = :comment_text
                    WHERE ID = :id
                    """,
                    {"id": step_id, "decision": d, "decided_by": (decided_by or "ui")[:120], "comment_text": (comment_text or "")[:2000] or None},
                )

                counts = _rows(
                    db.execute_query(
                        """
                        SELECT
                            SUM(CASE WHEN DECISION = 'PENDING' THEN 1 ELSE 0 END) AS PENDING_CNT,
                            SUM(CASE WHEN DECISION = 'REJECTED' THEN 1 ELSE 0 END) AS REJECTED_CNT,
                            SUM(CASE WHEN DECISION = 'APPROVED' THEN 1 ELSE 0 END) AS APPROVED_CNT
                        FROM CLS_CONTRACT_APPROVAL_STEPS
                        WHERE APPROVAL_ID = :id
                        """,
                        {"id": approval_id},
                    )
                )[0]

                approval = _rows(
                    db.execute_query(
                        "SELECT CONTRACT_ID FROM CLS_CONTRACT_APPROVALS WHERE ID = :id",
                        {"id": approval_id},
                    )
                )
                contract_id = int(approval[0]["contract_id"]) if approval else 0

                if int(counts.get("rejected_cnt") or 0) > 0:
                    db.execute_query(
                        "UPDATE CLS_CONTRACT_APPROVALS SET STATUS = 'REJECTED', FINISHED_AT = CURRENT_TIMESTAMP WHERE ID = :id",
                        {"id": approval_id},
                    )
                    db.execute_query(
                        "UPDATE CLS_CONTRACTS SET STATUS = 'REJECTED', UPDATED_AT = CURRENT_TIMESTAMP WHERE ID = :id",
                        {"id": contract_id},
                    )
                    final_status = "REJECTED"
                elif int(counts.get("pending_cnt") or 0) == 0:
                    db.execute_query(
                        "UPDATE CLS_CONTRACT_APPROVALS SET STATUS = 'APPROVED', FINISHED_AT = CURRENT_TIMESTAMP WHERE ID = :id",
                        {"id": approval_id},
                    )
                    db.execute_query(
                        "UPDATE CLS_CONTRACTS SET STATUS = 'APPROVED', UPDATED_AT = CURRENT_TIMESTAMP WHERE ID = :id",
                        {"id": contract_id},
                    )
                    final_status = "APPROVED"
                else:
                    final_status = "PENDING"

                db.connection.commit()
                return {"success": True, "approval_id": approval_id, "status": final_status}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def _required_attachment_flags(contract_id: int) -> Dict[str, int]:
        with DatabaseModel() as db:
            rows = _rows(
                db.execute_query(
                    """
                    SELECT HAS_FINANCIAL_TERMS, HAS_ESTIMATE_TERMS, HAS_PRICE_LIST
                    FROM V_CLS_CONTRACT_ATTACHMENT_CHECK
                    WHERE CONTRACT_ID = :id
                    """,
                    {"id": contract_id},
                )
            )
            if not rows:
                return {"financial": 0, "estimate": 0, "pricelist": 0}
            r = rows[0]
            return {
                "financial": int(r.get("has_financial_terms") or 0),
                "estimate": int(r.get("has_estimate_terms") or 0),
                "pricelist": int(r.get("has_price_list") or 0),
            }

    @staticmethod
    def export_contract_to_word(contract_id: int) -> Dict[str, Any]:
        try:
            if Document is None:
                return {"success": False, "error": "python-docx is not installed"}
            detail = ColassController.get_contract_detail(contract_id)
            if not detail.get("success"):
                return {"success": False, "error": detail.get("error") or "Contract not found"}
            data = detail.get("data") or {}
            attachments = ColassController.get_contract_attachments(contract_id)
            req = attachments.get("required") or {}
            if int(req.get("has_financial_terms") or 0) == 0:
                return {"success": False, "error": "Missing required attachment: FINANCIAL_TERMS"}
            if int(req.get("has_estimate_terms") or 0) == 0:
                return {"success": False, "error": "Missing required attachment: ESTIMATE_TERMS"}
            if int(req.get("has_price_list") or 0) == 0:
                return {"success": False, "error": "Missing required attachment: PRICE_LIST"}

            root = Path(__file__).resolve().parent.parent
            template_path = root / "docs" / "Colass" / "templates" / "contract_template.docx"
            if template_path.exists():
                doc = Document(str(template_path))
            else:
                doc = Document()

            doc.add_heading(f"Contract {data.get('contract_no') or ''}", level=1)
            doc.add_paragraph(f"Date: {(data.get('contract_date') or '')[:10]}")
            doc.add_paragraph(f"Company: {data.get('company_name') or ''}")
            doc.add_paragraph(f"Legal form: {data.get('legal_form') or ''}")
            doc.add_paragraph(f"Fiscal/VAT: {data.get('fiscal_code') or ''} / {data.get('vat_code') or ''}")
            doc.add_paragraph(f"Amount: {data.get('total_amount') or 0} {data.get('currency') or 'MDL'}")
            doc.add_paragraph(f"Status: {data.get('status') or ''}")
            doc.add_paragraph(f"Linked CRM lead: {data.get('lead_no') or ''}")
            doc.add_paragraph(f"Linked estimate: {data.get('estimate_name') or ''}")

            doc.add_heading("Contacts", level=2)
            emails = detail.get("emails") or []
            phones = detail.get("phones") or []
            routes = detail.get("routes") or []
            doc.add_paragraph("Emails: " + (", ".join([x.get("email") or "" for x in emails]) or "-"))
            doc.add_paragraph("Phones: " + (", ".join([x.get("phone") or "" for x in phones]) or "-"))
            doc.add_paragraph("Routes/Object: " + (", ".join([x.get("route_name") or "" for x in routes]) or "-"))

            doc.add_heading("Estimate conditions", level=2)
            doc.add_paragraph("Attachment type ESTIMATE_TERMS is provided and forms integral part of this contract.")
            doc.add_heading("Financial conditions", level=2)
            doc.add_paragraph("Attachment type FINANCIAL_TERMS is provided and forms integral part of this contract.")
            doc.add_heading("General price list", level=2)
            doc.add_paragraph("Attachment type PRICE_LIST is provided and forms integral part of this contract.")

            items = detail.get("items") or []
            doc.add_heading("Contract items", level=2)
            table = doc.add_table(rows=1, cols=6)
            hdr = table.rows[0].cells
            hdr[0].text = "Code"
            hdr[1].text = "Name"
            hdr[2].text = "Unit"
            hdr[3].text = "Qty"
            hdr[4].text = "Price"
            hdr[5].text = "Amount"
            for it in items:
                r = table.add_row().cells
                r[0].text = str(it.get("item_code") or "")
                r[1].text = str(it.get("item_name") or "")
                r[2].text = str(it.get("unit") or "")
                r[3].text = str(it.get("qty") or 0)
                r[4].text = str(it.get("price") or 0)
                r[5].text = str(it.get("amount") or 0)

            out = io.BytesIO()
            doc.save(out)
            out.seek(0)
            fname = (data.get("contract_no") or f"contract_{contract_id}").replace("/", "_") + ".docx"
            return {"success": True, "filename": fname, "content": out.getvalue(), "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    @staticmethod
    def import_crm_leads_from_email(limit: int = 20, unread_only: bool = True, created_by: str = "imap") -> Dict[str, Any]:
        if not Config.CRM_IMAP_ENABLED:
            return {"success": False, "error": "CRM_IMAP_ENABLED is false"}
        if not Config.CRM_IMAP_HOST or not Config.CRM_IMAP_USERNAME or not Config.CRM_IMAP_PASSWORD:
            return {"success": False, "error": "CRM IMAP settings are incomplete"}
        imported = 0
        deduplicated = 0
        errors: List[str] = []
        processed: List[Dict[str, Any]] = []

        conn = None
        try:
            if Config.CRM_IMAP_USE_SSL:
                conn = imaplib.IMAP4_SSL(Config.CRM_IMAP_HOST, Config.CRM_IMAP_PORT)
            else:
                conn = imaplib.IMAP4(Config.CRM_IMAP_HOST, Config.CRM_IMAP_PORT)
            conn.login(Config.CRM_IMAP_USERNAME, Config.CRM_IMAP_PASSWORD)
            conn.select(Config.CRM_IMAP_MAILBOX)

            criteria: List[str] = ["UNSEEN" if unread_only else "ALL"]
            if Config.CRM_IMAP_FROM_FILTER:
                criteria.append(f'FROM "{Config.CRM_IMAP_FROM_FILTER}"')
            typ, msgnums = conn.search(None, *criteria)
            if typ != "OK":
                return {"success": False, "error": "IMAP search failed"}
            ids = (msgnums[0].split() if msgnums and msgnums[0] else [])[-max(1, min(limit, 200)) :]

            for msg_id in ids:
                try:
                    typ_fetch, data = conn.fetch(msg_id, "(RFC822)")
                    if typ_fetch != "OK" or not data:
                        errors.append(f"fetch failed for {msg_id!r}")
                        continue
                    raw_msg = data[0][1] if isinstance(data[0], tuple) and len(data[0]) > 1 else None
                    if not raw_msg:
                        errors.append(f"empty payload for {msg_id!r}")
                        continue
                    msg = message_from_bytes(raw_msg)
                    message_id = (_decode_mime_value(msg.get("Message-ID") or "") or "").strip() or None
                    from_header = _decode_mime_value(msg.get("From") or "")
                    from_name, from_email = parseaddr(from_header)
                    subject = _decode_mime_value(msg.get("Subject") or "")
                    body = _extract_email_text(msg)
                    contact = _parse_contact_from_text((subject or "") + "\n" + (body or ""))
                    dt_obj: Optional[datetime] = None
                    try:
                        dt_obj = parsedate_to_datetime(msg.get("Date")) if msg.get("Date") else None
                    except Exception:
                        dt_obj = None

                    payload = {
                        "source_code": "EMAIL",
                        "stage_code": "NEW",
                        "contact_name": contact.get("contact_name") or from_name or None,
                        "company_name": contact.get("company_name"),
                        "phone": contact.get("phone"),
                        "email": contact.get("email") or from_email or None,
                        "subject": subject or None,
                        "needs_text": body[:3900] if body else None,
                        "lang_pref": _guess_lang_pref((subject or "") + " " + (body or "")),
                        "assigned_to": created_by,
                        "email_from": from_header or None,
                        "email_subject": subject or None,
                        "email_message_id": message_id,
                        "email_received_at": dt_obj,
                        "external_ref": message_id or None,
                    }
                    res = ColassController.create_crm_lead(payload, created_by=created_by)
                    if res.get("success") and res.get("deduplicated"):
                        deduplicated += 1
                    elif res.get("success"):
                        imported += 1
                        processed.append({"id": res.get("id"), "lead_no": res.get("lead_no"), "subject": subject})
                    else:
                        errors.append(res.get("error") or "create lead failed")
                except Exception as inner:
                    errors.append(str(inner))
            conn.logout()
            return {
                "success": True,
                "imported": imported,
                "deduplicated": deduplicated,
                "processed": processed,
                "errors": errors,
            }
        except Exception as e:
            try:
                if conn:
                    conn.logout()
            except Exception:
                pass
            return {"success": False, "error": str(e), "imported": imported, "deduplicated": deduplicated, "errors": errors}
